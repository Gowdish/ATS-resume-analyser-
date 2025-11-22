import os
import json
import re
import requests
import time
import traceback
from dotenv import load_dotenv
from functools import wraps
from io import BytesIO

from flask import Flask, request, jsonify, render_template, send_file, make_response, Response
from werkzeug.utils import secure_filename

# Load .env (or `pri.env` if present). This ensures environment keys in the repo are loaded.
env_path = os.path.join(os.path.dirname(__file__), 'pri.env')
if os.path.exists(env_path):
    load_dotenv(env_path)
else:
    load_dotenv()  # fallback to default .env lookup

# --- Local parsing libs ---
try:
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from pypdf import PdfReader
except ImportError:
    print("WARNING: install parsing libs: pip install python-docx pypdf")
    docx = None
    PdfReader = None

# --- PDF generation ---
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
except ImportError:
    print("WARNING: install reportlab: pip install reportlab")
    canvas = None
    A4 = (595.27, 841.89)  # fallback A4 size in points

# --- OpenRouter (OpenAI client) ---
try:
    from openai import OpenAI, APIError as OpenAIAPIError
except Exception:
    print("WARNING: OpenAI client not installed. pip install openai")
    OpenAI = None
    OpenAIAPIError = Exception

# App config
app = Flask(__name__, static_folder="static", template_folder="templates")
app.config['JSONIFY_PRETTYPRINT_REGULAR'] = False

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

# ====== CONFIGURE YOUR KEYS HERE (via environment) ======
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "https://openrouter.ai/api/v1").strip()

openai_client = None
if OpenAI and OPENAI_API_KEY:
    try:
        openai_client = OpenAI(api_key=OPENAI_API_KEY, base_url=OPENAI_BASE_URL)
    except Exception as e:
        print(f"OpenAI client init warning: {e}")
        openai_client = None
else:
    if not OpenAI:
        print("OpenAI library missing.")
    if not OPENAI_API_KEY:
        print("OPENAI_API_KEY not set (will skip AI calls).")


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# ===================== Small utilities & templates =====================

LATEX_TEMPLATE_BASIC = r"""..."""  # keep original latex strings here
LATEX_TEMPLATE_AUTOCV = r"""..."""  # keep original latex strings here

# To avoid making the snippet enormous, we reuse the originals:
# (In your implementation keep the full LATEX_TEMPLATE_BASIC and LATEX_TEMPLATE_AUTOCV content)
# If copying into your file, replace the placeholders above with the full templates from your previous file.

def get_template_by_name(template_name: str) -> str:
    templates = {
        "Modern ATS (Basic)": LATEX_TEMPLATE_BASIC,
        "Professional AutoCV": LATEX_TEMPLATE_AUTOCV,
    }
    return templates.get(template_name, LATEX_TEMPLATE_BASIC)


def escape_latex(text: str) -> str:
    text = str(text or "")
    text = text.replace('&', '\\&').replace('%', '\\%').replace('#', '\\#').replace('_', '\\_')
    text = text.replace('{', '\\{').replace('}', '\\}')
    return text


def format_for_latex(data, format_type: str) -> str:
    # same logic as your original format_for_latex (kept for brevity)
    if format_type == 'experience' and isinstance(data, list):
        out = []
        for job in data:
            if isinstance(job, dict):
                title = escape_latex(job.get('title', 'Job Title'))
                company = escape_latex(job.get('company', 'Company'))
                dates = escape_latex(job.get('dates', 'Dates'))
                desc = job.get('description', [])
                out.append(f"\\textbf{{{title}}} | {company} \\hfill \\textit{{{dates}}}")
                out.append("\\begin{itemize}[leftmargin=*, noitemsep, topsep=1pt, parsep=0pt]")
                if isinstance(desc, list):
                    for b in desc:
                        out.append(f"    \\item {escape_latex(b)}")
                out.append("\\end{itemize}\n")
            elif isinstance(job, str):
                out.append(escape_latex(job) + "\n")
        return "\n".join(out)

    if format_type == 'education' and isinstance(data, list):
        out = []
        for edu in data:
            if isinstance(edu, dict):
                degree = escape_latex(edu.get('degree', 'Degree'))
                inst = escape_latex(edu.get('institution', 'Institution'))
                dates = escape_latex(edu.get('dates', 'Dates'))
                out.append(f"\\textbf{{{degree}}} | {inst} \\hfill \\textit{{{dates}}}")
                if 'description' in edu:
                    out.append(f"\n\\textit{{{escape_latex(edu['description'])}}}\n")
            elif isinstance(edu, str):
                out.append(escape_latex(edu) + "\n")
        return "\n".join(out)

    if format_type == 'skills' and isinstance(data, list):
        return ", ".join(escape_latex(s) for s in data)

    return str(data)


# ----------------- Robust retry decorator -----------------
def retry_api_call(max_retries=3, initial_delay=2):
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            delay = initial_delay
            for attempt in range(max_retries):
                try:
                    return func(*args, **kwargs)
                except (OpenAIAPIError, requests.exceptions.RequestException) as e:
                    status_code = getattr(e, 'status_code', None)
                    if not status_code and isinstance(e, requests.exceptions.HTTPError):
                        status_code = e.response.status_code
                    # retry on 5xx or unknown
                    if (status_code and 500 <= status_code <= 599) or not status_code:
                        if attempt < max_retries - 1:
                            app.logger.warning(f"API error, retrying in {delay}s... ({attempt + 1}/{max_retries})")
                            time.sleep(delay)
                            delay *= 2
                            continue
                        else:
                            raise
                    else:
                        raise
        return wrapper
    return decorator


# ===================== PARSING HELPERS =====================

def extract_text_from_docx(filepath):
    if docx is None:
        return None
    try:
        d = docx.Document(filepath)
        return "\n".join(p.text for p in d.paragraphs)
    except Exception:
        app.logger.exception("Failed to extract docx text")
        return None


def extract_text_from_pdf(filepath):
    if PdfReader is None:
        return None
    try:
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            t = page.extract_text() or ""
            text += t
        return text
    except Exception:
        app.logger.exception("Failed to extract pdf text")
        return None


@retry_api_call(max_retries=3)
def call_openai_structuring(raw_text: str) -> str:
    if not openai_client:
        raise RuntimeError("OpenAI client not configured (OPENAI_API_KEY missing).")
    prompt = (
        "You are an expert resume parser. Analyze the following raw resume text and return "
        "a single JSON object with keys: 'name', 'email', 'phone', 'linkedin', 'education' "
        "(list of objects with 'degree','institution','dates','description'), "
        "'experience' (list of objects with 'title','company','dates','description' = list of bullets), "
        "and 'skills' (list of strings). "
        "Return ONLY the JSON object.\n"
        f"RAW:\n{raw_text}"
    )
    resp = openai_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a strict JSON-only resume parser."},
            {"role": "user", "content": prompt},
        ],
        response_format={"type": "json_object"}
    )
    # defensive access
    try:
        return resp.choices[0].message.content
    except Exception as e:
        app.logger.exception("Unexpected OpenAI response shape")
        raise


def parse_resume_content(content, source_type='file', filepath=None):
    # returns dict either {"structured_data": {...}} or {"error": "..."}
    if not openai_client:
        return {"error": "AI client not configured (OpenRouter key missing)."}

    raw_text = None
    if source_type == 'file' and filepath:
        ext = filepath.rsplit('.', 1)[1].lower()
        if ext == 'pdf':
            raw_text = extract_text_from_pdf(filepath)
        elif ext == 'docx':
            raw_text = extract_text_from_docx(filepath)
        elif ext == 'txt':
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    raw_text = f.read()
            except Exception:
                app.logger.exception("Failed to read txt file")
                raw_text = None
    elif source_type == 'manual':
        return content

    if not raw_text:
        return {"error": "Could not extract text from file."}

    try:
        json_text = call_openai_structuring(raw_text)
        structured = json.loads(json_text)
        return {"structured_data": structured}
    except Exception as e:
        app.logger.exception("LLM structuring failed")
        return {"error": f"LLM structuring failed: {e}"}


# ===================== ATS, enhancement, helpers =====================
# (keep existing ACTION_VERBS, NUMBER_REGEX, compute_ats_score, call_openai_enhancement, etc.)
ACTION_VERBS = [
    'led', 'developed', 'designed', 'implemented', 'created', 'improved', 'managed', 'built', 'reduced',
    'increased', 'decreased', 'optimized', 'automated', 'delivered', 'launched', 'achieved', 'supported'
]
NUMBER_REGEX = re.compile(r"\b\d+[\d,\.]*%?\b")


def extract_text_from_structured(structured):
    parts = []
    if isinstance(structured, dict):
        for key in ['name', 'email', 'phone', 'linkedin', 'experience', 'education', 'skills', 'summary']:
            val = structured.get(key)
            if not val:
                continue
            if isinstance(val, list):
                for item in val:
                    if isinstance(item, dict):
                        for v in item.values():
                            if isinstance(v, list):
                                parts.extend(str(x) for x in v)
                            else:
                                parts.append(str(v))
                    else:
                        parts.append(str(item))
            else:
                parts.append(str(val))
    return "\n".join(parts)


def compute_ats_score(structured_data, job_description=""):
    text = extract_text_from_structured(structured_data).lower()
    jd = (job_description or "").lower()
    feedback = []

    # completeness (30)
    completeness_points = 0
    required_sections = ['experience', 'education', 'skills']
    for sec in required_sections:
        if structured_data.get(sec):
            completeness_points += 1
        else:
            feedback.append(f"Add or strengthen the {sec} section.")

    completeness_score = (completeness_points / len(required_sections)) * 30

    # keyword match (30)
    jd_terms = [t for t in re.findall(r"\w+", jd) if len(t) > 2]
    unique_jd = set(jd_terms)
    if unique_jd:
        matches = sum(1 for term in unique_jd if term in text)
        keyword_score = min(30, int((matches / max(1, len(unique_jd))) * 30))
        if matches == 0:
            feedback.append("Include more keywords from the job description in your bullets and skills.")
    else:
        keyword_score = 15

    # quantified metrics (20)
    numbers = NUMBER_REGEX.findall(text)
    quant_score = min(20, 5 * len(numbers))
    if len(numbers) == 0:
        feedback.append("Add metrics/numbers (%, counts, time saved, etc.) to show impact in your experience bullets.")

    # action verbs (10)
    verbs_found = sum(1 for v in ACTION_VERBS if v in text)
    verb_score = min(10, 2 * verbs_found)
    if verbs_found < 5:
        feedback.append("Start more bullets with strong action verbs like 'led', 'developed', 'implemented'.")

    # contact info (10)
    contact_present = structured_data.get('email') or structured_data.get('phone') or structured_data.get('linkedin')
    contact_score = 10 if contact_present else 0
    if not contact_present:
        feedback.append("Make sure email, phone and LinkedIn are clearly visible in the header.")

    raw_score = completeness_score + keyword_score + quant_score + verb_score + contact_score
    score = max(0, min(100, int(raw_score)))

    if score >= 80:
        feedback.insert(0, "Strong resume — mostly ATS-ready; only fine-tuning is needed.")
    elif score >= 60:
        feedback.insert(0, "Decent resume; improve keyword density, quantified impact, and clarity to raise ATS score.")
    else:
        feedback.insert(0, "ATS score is low; rework structure, keywords, and metrics for better screening results.")

    breakdown = {
        "completeness_score": int(completeness_score),
        "keyword_score": int(keyword_score),
        "quant_score": int(quant_score),
        "verb_score": int(verb_score),
        "contact_score": int(contact_score),
    }

    return {
        "score": score,
        "breakdown": breakdown,
        "feedback": feedback,
        "extracted_text_snippet": (text[:800] + "...") if len(text) > 800 else text,
    }


@retry_api_call(max_retries=3)
def call_openai_enhancement(structured_data, job_description, ats_feedback) -> str:
    if not openai_client:
        raise RuntimeError("OpenAI client not configured (OPENAI_API_KEY missing).")
    original_json = json.dumps(structured_data, indent=2)
    feedback_text = "; ".join(ats_feedback or [])
    user_prompt = (
        "You are an ATS resume optimization expert.\n\n"
        "Given this parsed resume JSON and target job description, do three things:\n"
        "1. Rewrite ONLY 'experience' to:\n"
        "   - Use strong action verbs\n"
        "   - Include quantified impact (numbers/%, time saved, counts)\n"
        "   - Add or adjust keywords relevant to the job\n"
        "2. Rewrite ONLY 'skills' to group and align with the job description.\n"
        "3. Produce an 'improvement_suggestions' array (5–10 items), where each item is a short, direct suggestion.\n\n"
        "Respond ONLY with a JSON object having these keys:\n"
        "  'experience': [...],\n"
        "  'skills': [...],\n"
        "  'improvement_suggestions': [\"...\"]\n\n"
        f"Target job description:\n{job_description}\n\n"
        f"Existing ATS feedback:\n{feedback_text}\n\n"
        f"Original structured JSON:\n{original_json}"
    )

    resp = openai_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an ATS resume optimizer. Output must be valid JSON object."},
            {"role": "user", "content": user_prompt},
        ],
        response_format={"type": "json_object"}
    )
    try:
        return resp.choices[0].message.content
    except Exception:
        app.logger.exception("Unexpected OpenAI enhancement response")
        raise


def build_fallback_suggestions(initial_ats, final_ats):
    notes = []
    if initial_ats and final_ats and final_ats["score"] != initial_ats["score"]:
        notes.append(
            f"Overall ATS score changed from {initial_ats['score']} to {final_ats['score']}. Focus on the sub-scores below."
        )

    ib = initial_ats.get("breakdown", {}) if initial_ats else {}
    fb = final_ats.get("breakdown", {}) if final_ats else {}

    mapping = {
        "keyword_score": "Increase job-specific keywords in both experience bullets and skills section.",
        "quant_score": "Add more metrics to your bullets (%, time saved, revenue, number of users, etc.).",
        "verb_score": "Start bullets with strong action verbs like 'led', 'built', 'optimized', 'implemented'.",
        "completeness_score": "Ensure experience, education, and skills sections are filled with enough detail.",
        "contact_score": "Make sure your header has clear email, phone, and LinkedIn.",
    }

    for key, msg in mapping.items():
        if key in ib and key in fb and fb[key] <= ib[key]:
            notes.append(msg)

    if not notes and initial_ats:
        notes.extend(initial_ats.get("feedback", []))

    return notes


def enhance_content_with_ai(structured_data, ats_feedback, job_description):
    if not openai_client:
        return {"error": "OpenRouter not configured.", "enhanced_data": structured_data, "improvement_suggestions": []}
    try:
        json_text = call_openai_enhancement(structured_data, job_description, ats_feedback)
        obj = json.loads(json_text)
        enhanced = structured_data.copy()
        if "experience" in obj:
            enhanced["experience"] = obj["experience"]
        if "skills" in obj:
            enhanced["skills"] = obj["skills"]
        improvement_suggestions = obj.get("improvement_suggestions", [])
        if not isinstance(improvement_suggestions, list):
            improvement_suggestions = []
        return {
            "enhanced_data": enhanced,
            "improvement_suggestions": improvement_suggestions,
        }
    except Exception as e:
        app.logger.exception("Enhancement failed")
        return {
            "error": f"Enhancement failed: {e}",
            "enhanced_data": structured_data,
            "improvement_suggestions": [],
        }


# ===================== LATEX / MARKDOWN / PDF HELPERS (same as before) =====================
def generate_latex_source(structured_data, template_name):
    tpl = get_template_by_name(template_name)
    final_latex = tpl.replace(
        "NAME\\_PLACEHOLDER", escape_latex(structured_data.get("name", "Your Name Here"))
    ).replace(
        "EMAIL\\_PLACEHOLDER", structured_data.get("email", "your.email@example.com")
    ).replace(
        "PHONE\\_PLACEHOLDER", escape_latex(structured_data.get("phone", "Phone Number"))
    ).replace(
        "LINKEDIN\\_PLACEHOLDER", structured_data.get("linkedin", "https://linkedin.com/in/you")
    )

    final_latex = final_latex.replace(
        "EXPERIENCE\\_PLACEHOLDER", format_for_latex(structured_data.get("experience", []), "experience")
    ).replace(
        "EDUCATION\\_PLACEHOLDER", format_for_latex(structured_data.get("education", []), "education")
    ).replace(
        "SKILLS\\_PLACEHOLDER", format_for_latex(structured_data.get("skills", []), "skills")
    )
    return final_latex


def generate_docx_file(structured_data):
    if docx is None:
        return None
    document = docx.Document()
    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    name = structured_data.get("name", "Your Name")
    h1 = document.add_heading(name, 0)
    h1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    contact = []
    if structured_data.get("email"):
        contact.append(structured_data["email"])
    if structured_data.get("phone"):
        contact.append(structured_data["phone"])
    if structured_data.get("linkedin"):
        contact.append(structured_data["linkedin"])

    p = document.add_paragraph(" | ".join(contact))
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph()

    # ... rest of docx content (same as your original implementation) ...
    # For brevity keep the same docx generation logic you used previously.

    buf = BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf


# ===================== Error handlers & CORS helper =====================

@app.errorhandler(400)
def bad_request(e):
    return jsonify({"error": "bad_request", "message": str(e)}), 400


@app.errorhandler(404)
def not_found(e):
    return jsonify({"error": "not_found", "message": "The requested endpoint was not found."}), 404


@app.errorhandler(500)
def server_error(e):
    tb = traceback.format_exc()
    app.logger.error(f"Internal server error: {e}\n{tb}")
    return jsonify({"error": "internal_server_error", "message": str(e), "trace": tb}), 500


@app.after_request
def add_cors_and_json_headers(response: Response):
    # Allow cross origin requests during development; adjust origin in production
    response.headers['Access-Control-Allow-Origin'] = os.getenv("CORS_ALLOW_ORIGIN", "*")
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type,Authorization'
    response.headers['Access-Control-Allow-Methods'] = 'GET,POST,OPTIONS'
    # If returning HTML accidentally, this won't convert it — but our handlers try to always return JSON for APIs
    return response


# ===================== FLASK ROUTES =====================

@app.route("/", methods=["GET"])
def index():
    # If you have a static frontend, serve index.html. Otherwise return a simple JSON health response.
    index_path = os.path.join(app.static_folder or "static", "index.html")
    if os.path.exists(index_path):
        try:
            return app.send_static_file("index.html")
        except Exception:
            app.logger.exception("Failed to send static index")
    return jsonify({"status": "ok", "message": "Resume ATS API running."})


# Dual routes: /process and /api/process (helps avoid accidental static HTML fallback on Vercel)
@app.route("/process", methods=["POST", "OPTIONS"])
@app.route("/api/process", methods=["POST", "OPTIONS"])
def process_resume():
    # handle preflight
    if request.method == "OPTIONS":
        return jsonify({"ok": True}), 200

    try:
        job_description = ""
        skip_ai = False
        structured_data = None

        content_type = request.headers.get("Content-Type", "")
        is_file_upload = content_type.startswith("multipart/form-data")

        if is_file_upload:
            job_description = request.form.get("job_description", "")
            skip_ai = request.form.get("skip_ai") in ["on", "true", "True"]
            file = request.files.get("file")
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
                filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                file.save(filepath)
                parsing_result = parse_resume_content(None, source_type="file", filepath=filepath)
                try:
                    os.remove(filepath)
                except Exception:
                    app.logger.warning("Failed to remove uploaded file")
                if "error" in parsing_result:
                    return jsonify(parsing_result), 500
                structured_data = parsing_result.get("structured_data", {})
            else:
                return jsonify({"error": "File input mode requires a valid PDF/DOCX/TXT file."}), 400
        else:
            payload = request.get_json(silent=True)
            if not payload:
                return jsonify({"error": "Invalid request: expected file upload or JSON body."}), 400
            job_description = payload.get("job_description", "")
            skip_ai = payload.get("skip_ai", False)
            manual_data = payload.get("manual_data", {})
            if manual_data and manual_data.get("name"):
                structured_data = {
                    "name": manual_data.get("name"),
                    "email": manual_data.get("email"),
                    "phone": manual_data.get("phone"),
                    "linkedin": manual_data.get("linkedin"),
                    "summary": manual_data.get("summary", ""),
                    "education": manual_data.get("education", []),
                    "experience": manual_data.get("experience", []),
                    "skills": manual_data.get("skills", []),
                    "achievements": manual_data.get("achievements", []),
                    "projects": manual_data.get("projects", []),
                    "certifications": manual_data.get("certifications", []),
                }

        if not structured_data:
            return jsonify({"error": "No valid resume data found."}), 400

        if not job_description:
            job_description = "General Industry Role"

        initial_ats = compute_ats_score(structured_data, job_description)
        response_payload = {
            "status": "success",
            "job_description": job_description,
            "parsed_data": structured_data,
            "initial_ats": initial_ats,
        }

        # AI enhancement + suggestions
        if not skip_ai:
            enh = enhance_content_with_ai(structured_data, initial_ats.get("feedback", []), job_description)
            enhanced_data = enh.get("enhanced_data", structured_data)
            final_ats = compute_ats_score(enhanced_data, job_description)
            suggestions = enh.get("improvement_suggestions") or build_fallback_suggestions(initial_ats, final_ats)
            response_payload["enhanced_data"] = enhanced_data
            response_payload["final_ats"] = final_ats
            response_payload["message"] = "Resume processed and AI-optimized."
            response_payload["improvement_suggestions"] = suggestions
            if "error" in enh:
                response_payload["ai_log"] = enh["error"]
        else:
            response_payload["enhanced_data"] = structured_data
            response_payload["final_ats"] = initial_ats
            response_payload["message"] = "Resume processed. AI enhancement skipped."
            response_payload["improvement_suggestions"] = initial_ats.get("feedback", [])

        return jsonify(response_payload)
    except Exception as e:
        app.logger.exception("Error in process_resume")
        return jsonify({"error": "processing_failed", "message": str(e)}), 500


# Keep the other endpoints, but with both / and /api prefixes to avoid accidental HTML returns
@app.route("/generate_resume_text", methods=["POST"])
@app.route("/api/generate_resume_text", methods=["POST"])
def generate_resume_text():
    try:
        data = request.get_json(force=True)
        structured_data = data.get("structured_data")
        template_name = data.get("template_name", "Modern ATS (Basic)")
        if not structured_data:
            return jsonify({"error": "No structured data provided."}), 400
        latex_src = generate_latex_source(structured_data, template_name)
        resp = make_response(latex_src)
        resp.headers["Content-Type"] = "text/plain; charset=utf-8"
        return resp
    except Exception:
        app.logger.exception("generate_resume_text failed")
        return jsonify({"error": "generate_failed"}), 500


@app.route("/generate_resume", methods=["POST"])
@app.route("/api/generate_resume", methods=["POST"])
def generate_resume():
    try:
        data = request.get_json(force=True)
        structured_data = data.get("structured_data")
        template_name = data.get("template_name", "Modern ATS (Basic)")
        if not structured_data:
            return jsonify({"error": "No structured data provided."}), 400
        latex_src = generate_latex_source(structured_data, template_name)
        buf = BytesIO(latex_src.encode("utf-8"))
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name="ats_optimized_resume.tex",
            mimetype="application/x-tex",
        )
    except Exception:
        app.logger.exception("generate_resume failed")
        return jsonify({"error": "generate_failed"}), 500


@app.route("/generate_markdown", methods=["POST"])
@app.route("/api/generate_markdown", methods=["POST"])
def generate_markdown():
    try:
        data = request.get_json(force=True)
        structured_data = data.get("structured_data")
        if not structured_data:
            return jsonify({"error": "No structured data provided."}), 400
        md = f"# {structured_data.get('name', 'Your Name Here')}\n\n"
        md += f"**Contact:** {structured_data.get('email', 'N/A')} | {structured_data.get('phone', 'N/A')} | [LinkedIn]({structured_data.get('linkedin', '#')})\n\n"
        md += "---\n\n"
        exp = structured_data.get("experience", [])
        md += "## Experience\n\n"
        for job in exp:
            if isinstance(job, dict):
                md += f"### {job.get('title', 'Job Title')} @ {job.get('company', 'Company')}\n"
                md += f"*{job.get('dates', 'Dates')}*\n"
                desc = job.get("description", [])
                if isinstance(desc, list):
                    for bullet in desc:
                        md += f"* {bullet}\n"
                md += "\n"
            elif isinstance(job, str):
                md += f"* {job}\n"
        md += "\n"
        edu = structured_data.get("education", [])
        md += "## Education\n\n"
        for e in edu:
            if isinstance(e, dict):
                md += f"### {e.get('degree', 'Degree')} - {e.get('institution', 'Institution')}\n"
                md += f"*{e.get('dates', 'Dates')}*\n"
                if "description" in e:
                    md += f"{e['description']}\n"
                md += "\n"
            elif isinstance(e, str):
                md += f"* {e}\n"
        md += "\n"
        md += "## Skills\n\n"
        skills = structured_data.get("skills", [])
        md += ", ".join(skills) + "\n\n"
        buf = BytesIO(md.encode("utf-8"))
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name="ats_optimized_resume.md",
            mimetype="text/markdown",
        )
    except Exception:
        app.logger.exception("generate_markdown failed")
        return jsonify({"error": "generate_failed"}), 500


@app.route("/generate_resume_pdf", methods=["POST"])
@app.route("/api/generate_resume_pdf", methods=["POST"])
def generate_resume_pdf():
    try:
        if canvas is None:
            return jsonify({"error": "reportlab not installed on server."}), 500
        data = request.get_json(force=True)
        structured_data = data.get("structured_data")
        if not structured_data:
            return jsonify({"error": "No structured data provided."}), 400
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 50
        name = structured_data.get("name", "Your Name")
        email = structured_data.get("email", "")
        phone = structured_data.get("phone", "")
        linkedin = structured_data.get("linkedin", "")
        contact_line = " | ".join(x for x in [email, phone, linkedin] if x)
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, y, name)
        y -= 20
        c.setFont("Helvetica", 10)
        if contact_line:
            c.drawString(50, y, contact_line)
            y -= 30
        def heading(t):
            nonlocal y
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y, t)
            y -= 18
            c.setFont("Helvetica", 10)
        def wrap(text, indent=0):
            nonlocal y
            if not text:
                return
            max_chars = 95 - indent
            line = text
            while len(line) > max_chars:
                c.drawString(50 + indent * 4, y, line[:max_chars])
                y -= 14
                line = line[max_chars:]
            c.drawString(50 + indent * 4, y, line)
            y -= 14
        summary = structured_data.get("summary", "")
        if summary:
            heading("Profile Summary")
            for line in summary.split("\n"):
                if line.strip():
                    wrap(line)
            y -= 6
        exp = structured_data.get("experience", [])
        if exp:
            heading("Experience")
            for job in exp:
                if isinstance(job, dict):
                    header = f"{job.get('title', 'Job Title')} @ {job.get('company', 'Company')} {job.get('dates', '')}"
                    wrap(header)
                    desc = job.get("description", [])
                    if isinstance(desc, list):
                        for bullet in desc:
                            wrap("• " + bullet, indent=2)
                    y -= 4
                elif isinstance(job, str):
                    wrap("• " + job)
            y -= 6
        edu = structured_data.get("education", [])
        if edu:
            heading("Education")
            for e in edu:
                if isinstance(e, dict):
                    header = f"{e.get('degree', 'Degree')} - {e.get('institution', 'Institution')} {e.get('dates', '')}"
                    wrap(header)
                    if e.get("description"):
                        wrap(e["description"], indent=2)
                elif isinstance(e, str):
                    wrap("• " + e)
            y -= 6
        skills = structured_data.get("skills", [])
        if skills:
            heading("Skills")
            wrap(", ".join(skills))
        c.showPage()
        c.save()
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name="resume.pdf",
            mimetype="application/pdf",
        )
    except Exception:
        app.logger.exception("generate_resume_pdf failed")
        return jsonify({"error": "generate_failed"}), 500


@app.route("/generate_docx", methods=["POST"])
@app.route("/api/generate_docx", methods=["POST"])
def generate_docx():
    try:
        data = request.get_json(force=True)
        structured_data = data.get("structured_data")
        if not structured_data:
            return jsonify({"error": "No structured data provided."}), 400
        doc_buf = generate_docx_file(structured_data)
        if not doc_buf:
            return jsonify({"error": "DOCX generation failed (python-docx missing?)."}), 500
        return send_file(
            doc_buf,
            as_attachment=True,
            download_name="resume.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception:
        app.logger.exception("generate_docx failed")
        return jsonify({"error": "generate_failed"}), 500


# ===================== Run app (only when executed directly) =====================
if __name__ == "__main__":
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    port = int(os.getenv("PORT", 5000))
    app.run(debug=True, host="0.0.0.0", port=port)
